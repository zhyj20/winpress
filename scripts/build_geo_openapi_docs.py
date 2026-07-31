from __future__ import annotations

import re
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
MARKDOWN_PATH = DOCS / "WinPress-GEO客户开放API接口文档-2026-07-27.md"
DOCX_PATH = DOCS / "WinPress-GEO客户开放API接口文档-2026-07-27.docx"

FONT = "Microsoft YaHei"
INK = "172033"
MUTED = "5B6879"
BLUE = "174EA6"
BLUE_DARK = "0C3576"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
BORDER = "D9E2EF"
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_run_font(run, size: float, color: str = INK, bold: bool = False, italic: bool = False) -> None:
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def set_cell_margins(cell, top: int = 80, bottom: int = 80, start: int = 120, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        element = tc_mar.find(qn(f"w:{name}"))
        if element is None:
            element = OxmlElement(f"w:{name}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def shade(element, fill: str) -> None:
    properties = element._tc.get_or_add_tcPr() if hasattr(element, "_tc") else element._p.get_or_add_pPr()
    node = properties.find(qn("w:shd"))
    if node is None:
        node = OxmlElement("w:shd")
        properties.append(node)
    node.set(qn("w:fill"), fill)


def add_paragraph_border(paragraph, side: str, color: str, size: str = "12", space: str = "8") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    border = OxmlElement(f"w:{side}")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), size)
    border.set(qn("w:space"), space)
    border.set(qn("w:color"), color)
    borders.append(border)


def configure_document(doc: Document) -> None:
    # LibreOffice honors the even-page header part only when this flag is explicit.
    # Both parts carry identical content, so every page keeps the same running label.
    doc.settings.odd_and_even_pages_header_footer = True
    section = doc.sections[0]
    section.different_first_page_header_footer = False
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, BLUE_DARK, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, BLUE_DARK, 10, 5),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = styles[name]
        style.font.name = FONT
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = FONT
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    if "Code Block" not in styles:
        code = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code = styles["Code Block"]
    code.font.name = "Consolas"
    code._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Consolas")
    code._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Consolas")
    code._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    code.font.size = Pt(8.5)
    code.font.color.rgb = RGBColor.from_string(INK)
    code.paragraph_format.left_indent = Inches(0.15)
    code.paragraph_format.right_indent = Inches(0.15)
    code.paragraph_format.space_before = Pt(4)
    code.paragraph_format.space_after = Pt(8)
    code.paragraph_format.line_spacing = 1.1

    for header_part in (section.header, section.even_page_header, section.first_page_header):
        header_part.is_linked_to_previous = False
        header = header_part.paragraphs[0]
        header.text = "WINPRESS 云发布 · GEO 服务端联动接口"
        header.alignment = WD_ALIGN_PARAGRAPH.LEFT
        header.paragraph_format.space_after = Pt(2)
        set_run_font(header.runs[0], 8.5, MUTED)
        add_paragraph_border(header, "bottom", BORDER, "6", "4")

    for footer_part in (section.footer, section.even_page_footer, section.first_page_footer):
        footer_part.is_linked_to_previous = False
        footer = footer_part.paragraphs[0]
        footer.text = ""
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = footer.add_run("实现核对版  |  ")
        set_run_font(run, 8.5, MUTED)
        page_run = footer.add_run()
        set_run_font(page_run, 8.5, MUTED)
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        instruction = OxmlElement("w:instrText")
        instruction.set(qn("xml:space"), "preserve")
        instruction.text = " PAGE "
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        page_run._r.extend((begin, instruction, end))


def add_title_block(doc: Document) -> None:
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_before = Pt(14)
    kicker.paragraph_format.space_after = Pt(4)
    set_run_font(kicker.add_run("WINPRESS · API REFERENCE"), 9, BLUE, True)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(6)
    set_run_font(title.add_run("GEO 服务端联动接口"), 24, BLUE_DARK, True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(16)
    set_run_font(subtitle.add_run("按当前源码核对的签名、报价、渠道目录与订单受理契约"), 11.5, MUTED)

    meta = doc.add_table(rows=2, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.LEFT
    meta.autofit = False
    rows = (("版本", "v1.1 实现核对版"), ("状态", "本机可验证；生产配置与联合验收待确认"))
    for row_index, values in enumerate(rows):
        for column_index, value in enumerate(values):
            cell = meta.cell(row_index, column_index)
            cell.text = ""
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, 90, 90, 120, 120)
            shade(cell, LIGHT_BLUE if column_index == 0 else "FFFFFF")
            paragraph = cell.paragraphs[0]
            set_run_font(paragraph.add_run(value), 9, BLUE_DARK if column_index == 0 else INK, column_index == 0)
    set_table_geometry(meta, [1500, CONTENT_WIDTH_DXA - 1500])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def parse_inline(paragraph, text: str) -> None:
    position = 0
    pattern = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*)")
    for match in pattern.finditer(text):
        if match.start() > position:
            set_run_font(paragraph.add_run(text[position:match.start()]), 10.5)
        token = match.group(0)
        if token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, 9.2, BLUE_DARK)
            run.font.name = "Consolas"
        else:
            set_run_font(paragraph.add_run(token[2:-2]), 10.5, INK, True)
        position = match.end()
    if position < len(text):
        set_run_font(paragraph.add_run(text[position:]), 10.5)


def set_table_geometry(table, widths: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table_pr = table._tbl.tblPr
    layout = table_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        table_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    width = table_pr.find(qn("w:tblW"))
    if width is None:
        width = OxmlElement("w:tblW")
        table_pr.append(width)
    width.set(qn("w:type"), "dxa")
    width.set(qn("w:w"), str(sum(widths)))
    indent = table_pr.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd")
        table_pr.append(indent)
    indent.set(qn("w:type"), "dxa")
    indent.set(qn("w:w"), str(TABLE_INDENT_DXA))

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for column_width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(column_width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_width = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_width is None:
                tc_width = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_width)
            tc_width.set(qn("w:type"), "dxa")
            tc_width.set(qn("w:w"), str(widths[index]))


def table_widths(rows: list[list[str]]) -> list[int]:
    columns = len(rows[0])
    weights = []
    for index in range(columns):
        longest = max(len(row[index]) for row in rows)
        weights.append(max(8, min(longest, 48)))
    total = sum(weights)
    widths = [max(900, round(CONTENT_WIDTH_DXA * weight / total)) for weight in weights]
    difference = CONTENT_WIDTH_DXA - sum(widths)
    widths[-1] += difference
    if widths[-1] < 900:
        shortage = 900 - widths[-1]
        widths[-1] = 900
        widest = max(range(len(widths) - 1), key=widths.__getitem__)
        widths[widest] -= shortage
    return widths


def clean_table_text(value: str) -> str:
    return value.replace("`", "").replace("**", "").strip()


def add_table(doc: Document, rows: list[list[str]]) -> None:
    rows = [[clean_table_text(value) for value in row] for row in rows]
    # Named layout override: keep the four-service mapping together instead of leaving an
    # isolated header row at the bottom of the preceding page.
    if rows[0][:2] == ["服务", "service_type"]:
        doc.add_page_break()
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for row_index, values in enumerate(rows):
        for column_index, value in enumerate(values):
            cell = table.cell(row_index, column_index)
            cell.text = ""
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if row_index == 0:
                shade(cell, LIGHT_BLUE)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            if row_index == 0:
                paragraph.paragraph_format.keep_with_next = True
            set_run_font(
                paragraph.add_run(value),
                8.6 if row_index else 8.8,
                BLUE_DARK if row_index == 0 else INK,
                row_index == 0,
            )
    header_properties = table.rows[0]._tr.get_or_add_trPr()
    repeat_header = OxmlElement("w:tblHeader")
    repeat_header.set(qn("w:val"), "true")
    header_properties.append(repeat_header)
    set_table_geometry(table, table_widths(rows))
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_code_block(doc: Document, lines: Iterable[str]) -> None:
    paragraph = doc.add_paragraph(style="Code Block")
    shade(paragraph, LIGHT_GRAY)
    add_paragraph_border(paragraph, "left", BLUE, "16", "6")
    for index, line in enumerate(lines):
        if index:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line)
        run.font.name = "Consolas"
        run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Consolas")
        run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Consolas")
        run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
        run.font.size = Pt(8.3)
        run.font.color.rgb = RGBColor.from_string(INK)


def build_docx(markdown: str) -> None:
    doc = Document()
    configure_document(doc)
    add_title_block(doc)

    lines = markdown.splitlines()
    index = 0
    paragraph_buffer: list[str] = []

    def flush_paragraph() -> None:
        if not paragraph_buffer:
            return
        paragraph = doc.add_paragraph()
        parse_inline(paragraph, " ".join(item.strip() for item in paragraph_buffer))
        paragraph_buffer.clear()

    while index < len(lines):
        raw = lines[index]
        line = raw.strip()
        if not line:
            flush_paragraph()
            index += 1
            continue
        if line.startswith("```"):
            flush_paragraph()
            index += 1
            code_lines: list[str] = []
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code_lines.append(lines[index])
                index += 1
            add_code_block(doc, code_lines)
            index += 1
            continue
        if line.startswith("|") and index + 1 < len(lines) and re.match(r"^\|[| :\-]+\|$", lines[index + 1].strip()):
            flush_paragraph()
            table_rows = [[cell.strip() for cell in line.strip("|").split("|")]]
            index += 2
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_rows.append([cell.strip() for cell in lines[index].strip().strip("|").split("|")])
                index += 1
            add_table(doc, table_rows)
            continue
        heading = re.match(r"^(#{1,3})\s+(.+)$", line)
        if heading:
            flush_paragraph()
            text = heading.group(2)
            if text == "WinPress 云发布 GEO 服务端联动接口":
                index += 1
                continue
            # Named layout override: keep the production activation checklist together.
            if text == "10. 部署与验收":
                doc.add_page_break()
            doc.add_paragraph(text, style=f"Heading {len(heading.group(1))}")
            index += 1
            continue
        if line.startswith(">"):
            flush_paragraph()
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.15)
            paragraph.paragraph_format.right_indent = Inches(0.1)
            paragraph.paragraph_format.space_before = Pt(4)
            paragraph.paragraph_format.space_after = Pt(10)
            shade(paragraph, LIGHT_BLUE)
            add_paragraph_border(paragraph, "left", BLUE, "18", "8")
            parse_inline(paragraph, line[1:].strip())
            index += 1
            continue
        if re.match(r"^-\s+", line):
            flush_paragraph()
            paragraph = doc.add_paragraph(style="List Bullet")
            parse_inline(paragraph, re.sub(r"^-\s+", "", line))
            index += 1
            continue
        if re.match(r"^\d+\.\s+", line):
            flush_paragraph()
            paragraph = doc.add_paragraph(style="List Number")
            parse_inline(paragraph, re.sub(r"^\d+\.\s+", "", line))
            index += 1
            continue
        if re.match(r"^(版本|核对日期|适用对象)：", line):
            flush_paragraph()
            index += 1
            continue
        paragraph_buffer.append(line)
        index += 1

    flush_paragraph()
    final_section = doc.sections[-1]
    final_section.start_type = WD_SECTION.NEW_PAGE if len(doc.sections) > 1 else final_section.start_type
    doc.core_properties.title = "WinPress 云发布 GEO 服务端联动接口"
    doc.core_properties.subject = "按当前源码核对的服务端联动契约"
    doc.core_properties.author = "WinPress"
    doc.core_properties.keywords = "WinPress,GEO,API,server federation"
    doc.save(DOCX_PATH)


def main() -> None:
    markdown = MARKDOWN_PATH.read_text(encoding="utf-8")
    build_docx(markdown)
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
